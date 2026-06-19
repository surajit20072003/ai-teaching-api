"""
core/document_processor.py
───────────────────────────
Handles PDF/DOCX/TXT document ingestion:
  1. extract_text()   — pulls raw text from uploaded file
  2. chunk_text()     — splits into ~400-word overlapping passages
  3. embed_chunks()   — generates 384-dim vectors via sentence-transformers
  4. save_locally()   — writes raw file to /sdb-disk
  5. save_processed() — writes extracted.txt to /sdb-disk
"""

import re
import logging
from typing import Optional
from pathlib import Path

logger = logging.getLogger(__name__)

# ── Text Extraction ────────────────────────────────────────────────────────────

def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract plain text from PDF, DOCX, or TXT file bytes.
    Returns the full text as a single string.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_bytes)
    elif ext == ".txt":
        return file_bytes.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use PDF, DOCX, or TXT.")


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber (handles tables and multi-column)."""
    import pdfplumber
    import io

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
            if page_text and page_text.strip():
                text_parts.append(f"\n--- Page {page_num} ---\n{page_text.strip()}")

    full_text = "\n".join(text_parts)
    logger.info(f"[doc_processor] PDF extracted: {len(full_text):,} chars, {len(pdf.pages)} pages")
    return full_text


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    import docx
    import io

    doc = docx.Document(io.BytesIO(file_bytes))
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            # Detect headings by style name
            if para.style.name.startswith("Heading"):
                text_parts.append(f"\n## {para.text.strip()}\n")
            else:
                text_parts.append(para.text.strip())

    full_text = "\n".join(text_parts)
    logger.info(f"[doc_processor] DOCX extracted: {len(full_text):,} chars, {len(doc.paragraphs)} paragraphs")
    return full_text


# ── Text Chunking ──────────────────────────────────────────────────────────────

def chunk_text(
    text: str,
    chunk_size: int = 400,      # target words per chunk
    overlap: int = 50,          # words of overlap between chunks
) -> list[dict]:
    """
    Split extracted text into overlapping chunks.

    Returns a list of dicts:
    [
      {
        "chunk_index": 0,
        "section_title": "Chapter 3: Forces",  # or None
        "chunk_text": "Newton's second law states ...",
        "word_count": 397
      },
      ...
    ]
    """
    # Detect section headings (## Title or Page N lines)
    heading_pattern = re.compile(
        r'^(?:## (.+)|--- Page \d+ ---|([A-Z][A-Z\s]{3,50}):?\s*$)',
        re.MULTILINE
    )

    # Split into paragraphs for natural chunk boundaries
    paragraphs = [p.strip() for p in re.split(r'\n{2,}', text) if p.strip()]

    chunks = []
    current_words = []
    current_title: Optional[str] = None
    chunk_index = 0

    for para in paragraphs:
        # Check if paragraph is a section heading
        heading_match = heading_pattern.match(para)
        if heading_match:
            # Save current chunk before starting new section
            if current_words and len(current_words) >= 20:
                chunks.append(_make_chunk(chunk_index, current_title, current_words))
                chunk_index += 1
                # Keep last `overlap` words into next chunk
                current_words = current_words[-overlap:]

            current_title = (heading_match.group(1) or para).strip()
            continue

        para_words = para.split()

        # If adding this para exceeds chunk_size, flush first
        if len(current_words) + len(para_words) > chunk_size and current_words:
            chunks.append(_make_chunk(chunk_index, current_title, current_words))
            chunk_index += 1
            # Start next chunk with overlap from end of previous
            current_words = current_words[-overlap:]

        current_words.extend(para_words)

    # Flush remaining words
    if current_words and len(current_words) >= 20:
        chunks.append(_make_chunk(chunk_index, current_title, current_words))

    logger.info(f"[doc_processor] Chunked into {len(chunks)} chunks (target={chunk_size}w, overlap={overlap}w)")
    return chunks


def _make_chunk(index: int, title: Optional[str], words: list[str]) -> dict:
    return {
        "chunk_index": index,
        "section_title": title,
        "chunk_text": " ".join(words),
        "word_count": len(words),
    }


# ── Embedding ─────────────────────────────────────────────────────────────────

_model = None  # lazy-loaded singleton

def get_embedding_model():
    """Lazy-load sentence-transformers model (only once per process)."""
    global _model
    if _model is None:
        from sentence_transformers import SentenceTransformer
        logger.info("[doc_processor] Loading embedding model all-MiniLM-L6-v2 ...")
        _model = SentenceTransformer("all-MiniLM-L6-v2")
        logger.info("[doc_processor] Embedding model loaded")
    return _model


def embed_chunks(chunks: list[dict]) -> list[dict]:
    """
    Add 'embedding' (list[float], 384-dim) to each chunk dict.
    Batch processes all chunks in one shot for efficiency.
    Returns the same list with embeddings added.
    """
    if not chunks:
        return chunks

    model = get_embedding_model()
    texts = [c["chunk_text"] for c in chunks]

    # Batch embed — sentence-transformers handles batching internally
    embeddings = model.encode(texts, batch_size=32, show_progress_bar=False)

    for chunk, emb in zip(chunks, embeddings):
        chunk["embedding"] = emb.tolist()  # convert numpy -> plain list for JSON/pgvector

    logger.info(f"[doc_processor] Embedded {len(chunks)} chunks")
    return chunks


def embed_single(text: str) -> list[float]:
    """Embed a single query string. Used at search time."""
    model = get_embedding_model()
    return model.encode([text], show_progress_bar=False)[0].tolist()


# ── Local File Save (wrappers around local_storage) ──────────────────────────

async def save_locally(
    subject_id: str,
    doc_id: str,
    filename: str,
    file_bytes: bytes,
) -> str:
    """Save raw uploaded file to /sdb-disk and return local path."""
    from core.local_storage import write_raw_file, ensure_subject_dirs
    ensure_subject_dirs(subject_id)
    return await write_raw_file(subject_id, doc_id, filename, file_bytes)


async def save_processed(subject_id: str, doc_id: str, text: str) -> str:
    """Save extracted plain text to /sdb-disk and return local path."""
    from core.local_storage import write_processed_text
    return await write_processed_text(subject_id, doc_id, text)


# ── Full Pipeline (convenience function) ─────────────────────────────────────

async def process_document(
    subject_id: str,
    doc_id: str,
    filename: str,
    file_bytes: bytes,
) -> dict:
    """
    Run the full ingestion pipeline:
      1. Save raw file locally
      2. Extract text
      3. Save extracted text locally
      4. Chunk text
      5. Embed chunks

    Returns:
    {
      "local_raw_path": "/sdb-disk/...",
      "local_processed_path": "/sdb-disk/...",
      "chunks": [...],       # list of chunk dicts with embeddings
      "total_chars": 12345,
      "total_chunks": 32,
    }
    """
    # Step 1: Save raw file
    local_raw = await save_locally(subject_id, doc_id, filename, file_bytes)

    # Step 2: Extract text
    text = extract_text(file_bytes, filename)
    if not text or not text.strip():
        raise ValueError(f"Could not extract any text from {filename}")

    # Step 3: Save extracted text
    local_processed = await save_processed(subject_id, doc_id, text)

    # Step 4: Chunk
    chunks = chunk_text(text)

    # Step 5: Embed (CPU — runs in ~1-5s for typical lecture doc)
    chunks = embed_chunks(chunks)

    logger.info(
        f"[doc_processor] Document processed: subject={subject_id} doc={doc_id} "
        f"chars={len(text):,} chunks={len(chunks)}"
    )

    return {
        "local_raw_path": local_raw,
        "local_processed_path": local_processed,
        "chunks": chunks,
        "total_chars": len(text),
        "total_chunks": len(chunks),
    }
